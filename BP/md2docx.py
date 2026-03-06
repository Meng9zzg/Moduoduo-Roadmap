"""将 BP初稿4.md 转换为格式化的 Word 文档"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    from lxml import etree
    tc_pr = cell._element.get_or_add_tcPr()
    shading_el = etree.SubElement(tc_pr, qn('w:shd'))
    shading_el.set(qn('w:fill'), color_hex)
    shading_el.set(qn('w:val'), 'clear')

def create_doc():
    with open(r"E:\Jzzg\源码及技术路线\Moduoduo Roadmap\BP\BP初稿4.md", "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = '微软雅黑'
        hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row_data in enumerate(table_rows):
            for ci, cell_text in enumerate(row_data):
                cell = table.cell(ri, ci)
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(10)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if ri == 0:
                        for run in paragraph.runs:
                            run.bold = True
                        set_cell_shading(cell, 'E8EAF6')
        doc.add_paragraph()
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```') and not in_code_block:
            flush_table()
            in_code_block = True
            code_lines = []
            i += 1
            continue
        elif line.strip().startswith('```') and in_code_block:
            in_code_block = False
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_lines = []
            i += 1
            continue
        elif in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            cols = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(re.match(r'^[-:]+$', c) for c in cols):
                i += 1
                continue
            if not in_table:
                flush_table()
                in_table = True
            table_rows.append(cols)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            if text.startswith('Slide'):
                doc.add_page_break()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.strip() == '---':
            i += 1
            continue

        if line.strip().startswith('> '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        if line.strip().startswith('- '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        if line.strip() == '':
            i += 1
            continue

        text = line.strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        i += 1

    flush_table()

    output = r"E:\Jzzg\源码及技术路线\Moduoduo Roadmap\BP\BP初稿4.docx"
    doc.save(output)
    print(f"Word 文档已生成：{output}")

if __name__ == '__main__':
    create_doc()
